"""
AI_Gate Document Processor - document_scraper.py

A sophisticated document crawler and content processor that generates structured JSON data 
from target website documents (PDF, DOCX, XLSX) for integration with AI_Gate intelligent systems.

This version is fully compatible with website_researcher.py requirements and follows
the same structure as website_scraper.py.

REFACTORED VERSION: Now uses centralized AdvancedContentAnalyzer for consistent, 
high-quality multilingual analysis across all content sources.
"""

import json
import os
import time
import re
import hashlib
import sys

# --- Path Correction ---
# Add the project's root directory to the Python path.
# This allows this script to be run from anywhere and still import modules 
# from the 'modules' directory correctly, just as if it were run from the root.
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
import requests
from typing import Dict, List, Tuple, Optional
from datetime import datetime
import logging
from http.client import HTTPConnection

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import openpyxl
from openpyxl.utils import get_column_letter
import io

# Central analyzer import - the only NLP logic connection
from modules.analyzer import AdvancedContentAnalyzer

# Debug configuration
DEBUG = False
if DEBUG:
    HTTPConnection.debuglevel = 1
    logging.basicConfig()
    requests_log = logging.getLogger("requests.packages.urllib3")
    requests_log.setLevel(logging.DEBUG)
    requests_log.propagate = True

# Directories (updated paths for project structure)
CONFIG_DIR = "./config"
OUTPUT_DIR = "./data"  # Changed from ./data/document_data to ./data
LOG_DIR = "./logs"
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(LOG_DIR, exist_ok=True)

# Logging
log_file = os.path.join(LOG_DIR, f"document_processor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(name)s - %(message)s",
    handlers=[logging.FileHandler(log_file, encoding='utf-8'), logging.StreamHandler()]
)
logger = logging.getLogger("AI_Gate_Document_Processor")

class PDFProcessor:
    """Handle PDF document processing."""
    
    @staticmethod
    def extract_content(file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """Extract text content and metadata from PDF."""
        try:
            # Try pdfplumber first (better text extraction)
            pdf_file = io.BytesIO(file_content)
            with pdfplumber.open(pdf_file) as pdf:
                text_content = ""
                structured_content = {
                    'total_pages': len(pdf.pages),
                    'pages': []
                }
                
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                        
                        # Store first few pages for structure
                        if i < 5:
                            structured_content['pages'].append({
                                'page_number': i + 1,
                                'text': page_text[:1000] + "..." if len(page_text) > 1000 else page_text,
                                'has_tables': bool(page.extract_tables()),
                                'char_count': len(page_text)
                            })
                
                # If we got content, return it
                if text_content.strip():
                    return text_content.strip(), structured_content
                    
        except Exception as e:
            logger.warning(f"pdfplumber failed for {filename}: {str(e)}")
        
        # Fallback to PyPDF2
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            structured_content = {
                'total_pages': len(pdf_reader.pages),
                'pages': []
            }
            
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text_content += page_text + "\n"
                    
                    # Store first few pages for structure
                    if i < 5:
                        structured_content['pages'].append({
                            'page_number': i + 1,
                            'text': page_text[:1000] + "..." if len(page_text) > 1000 else page_text,
                            'char_count': len(page_text)
                        })
                        
                except Exception as page_error:
                    logger.warning(f"Error extracting page {i+1} from {filename}: {str(page_error)}")
                    continue
            
            # Check if we extracted meaningful content
            if text_content.strip() and len(text_content.strip()) > 50:
                return text_content.strip(), structured_content
            else:
                logger.warning(f"PDF {filename} appears to be scanned or contains no extractable text")
                return "", {}
                
        except Exception as e:
            logger.error(f"Both PDF processors failed for {filename}: {str(e)}")
            return "", {}

class DOCXProcessor:
    """Handle DOCX document processing."""
    
    @staticmethod
    def extract_content(file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """Extract text content and metadata from DOCX."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_content = ""
            structured_content = {
                'paragraphs': [],
                'headings': [],
                'tables': []
            }
            
            # Extract paragraphs and identify headings
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    text_content += para_text + "\n"
                    
                    # Check if paragraph is a heading
                    if paragraph.style.name.startswith('Heading'):
                        structured_content['headings'].append({
                            'level': int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1,
                            'text': para_text
                        })
                    else:
                        # Store first few paragraphs
                        if len(structured_content['paragraphs']) < 10:
                            structured_content['paragraphs'].append(para_text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(cell for cell in row_data):  # Skip empty rows
                        table_data.append(row_data)
                
                if table_data and len(structured_content['tables']) < 3:  # Limit tables
                    structured_content['tables'].append({
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0,
                        'sample_data': table_data[:3]  # First 3 rows
                    })
            
            return text_content.strip(), structured_content
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {filename}: {str(e)}")
            return "", {}

class XLSXProcessor:
    """Handle XLSX document processing."""
    
    @staticmethod
    def extract_content(file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """Extract text content and metadata from XLSX."""
        try:
            xlsx_file = io.BytesIO(file_content)
            workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
            
            text_content = ""
            structured_content = {
                'worksheets': [],
                'total_sheets': len(workbook.worksheets)
            }
            
            for sheet in workbook.worksheets:
                sheet_data = {
                    'name': sheet.title,
                    'dimensions': f"{sheet.max_row}x{sheet.max_column}",
                    'sample_data': [],
                    'headers': []
                }
                
                # Extract text content from cells
                sheet_text = f"Sheet: {sheet.title}\n"
                
                # Get headers (first row)
                header_row = []
                for col in range(1, min(sheet.max_column + 1, 11)):  # Limit to 10 columns
                    cell_value = sheet.cell(row=1, column=col).value
                    if cell_value is not None:
                        header_row.append(str(cell_value))
                    else:
                        header_row.append("")
                
                if any(header_row):
                    sheet_data['headers'] = header_row
                    sheet_text += "Headers: " + " | ".join(header_row) + "\n"
                
                # Sample data (first few rows)
                for row in range(1, min(sheet.max_row + 1, 6)):  # First 5 rows
                    row_data = []
                    for col in range(1, min(sheet.max_column + 1, 11)):  # First 10 columns
                        cell_value = sheet.cell(row=row, column=col).value
                        if cell_value is not None:
                            row_data.append(str(cell_value))
                            # Add to text content for searching
                            if isinstance(cell_value, str) and len(cell_value) > 3:
                                sheet_text += str(cell_value) + " "
                        else:
                            row_data.append("")
                    
                    if any(row_data):
                        sheet_data['sample_data'].append(row_data)
                
                text_content += sheet_text + "\n\n"
                structured_content['worksheets'].append(sheet_data)
            
            return text_content.strip(), structured_content
            
        except Exception as e:
            logger.error(f"Failed to process XLSX {filename}: {str(e)}")
            return "", {}

class DocumentProcessor:
    def __init__(self, config_path: str = os.path.join(CONFIG_DIR, "sites.json")):
        self.config = self._load_config(config_path)
        self.base_url = self._normalize_url(self.config["website_url"])
        self.session = self._configure_session()
        self.start_time = time.time()
        self.documents_data = []  # Store all document data for consolidated output
        
        # Initialize central content analyzer - the only NLP connection
        self.content_analyzer = AdvancedContentAnalyzer()
        
        # Initialize document processors (keep extraction helpers)
        self.pdf_processor = PDFProcessor()
        self.docx_processor = DOCXProcessor()
        self.xlsx_processor = XLSXProcessor()
        
        logger.info(f"Initializing document processor for {self.base_url}")

    def _load_config(self, config_path: str) -> Dict:
        """Load configuration from JSON file."""
        if not os.path.exists(config_path):
            raise FileNotFoundError(f"Configuration file not found: {config_path}")
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _normalize_url(self, url: str) -> str:
        """Normalize URL format."""
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        return url.rstrip('/')

    def _configure_session(self) -> requests.Session:
        """Configure HTTP session with headers and timeouts."""
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'AI_Gate_DocumentCrawler/2.0 (+https://ai-gate.org/bot)',
            'Accept-Language': 'en-US,en;q=0.9,ar;q=0.8',
            'Accept': '*/*'
        })
        session.max_redirects = 3
        session.timeout = 30  # Longer timeout for documents
        return session

    def _verify_site_accessibility(self) -> Tuple[bool, str]:
        """Verify that the target website is accessible."""
        try:
            logger.info("Verifying website accessibility...")
            home_resp = self.session.get(self.base_url, timeout=15)
            home_resp.raise_for_status()
            return True, "Website is accessible"
        except requests.exceptions.SSLError:
            logger.warning("SSL verification failed - retrying without verification")
            self.session.verify = False
            try:
                test_resp = self.session.get(self.base_url, timeout=10)
                test_resp.raise_for_status()
                return True, "Accessible (SSL verification disabled)"
            except Exception as e:
                return False, f"SSL failed: {str(e)}"
        except Exception as e:
            return False, f"Access check failed: {str(e)}"

    def _is_document_url(self, url: str) -> bool:
        """Check if URL points to a supported document type."""
        url_lower = url.lower()
        return any(url_lower.endswith(ext) for ext in ['.pdf', '.docx', '.xlsx'])

    def _get_document_type(self, url: str) -> str:
        """Get document type from URL."""
        url_lower = url.lower()
        if url_lower.endswith('.pdf'):
            return 'pdf'
        elif url_lower.endswith('.docx'):
            return 'docx'
        elif url_lower.endswith('.xlsx'):
            return 'xlsx'
        return 'unknown'

    def _discover_document_urls(self) -> List[str]:
        """Discover document URLs from the website."""
        document_urls = set()
        
        # First, check sitemap for documents
        sitemap_urls = self._crawl_sitemap_for_documents()
        document_urls.update(sitemap_urls)
        
        # Then crawl pages to find document links
        page_urls = self._discover_documents_from_pages()
        document_urls.update(page_urls)
        
        logger.info(f"Discovered {len(document_urls)} document URLs")
        return sorted(document_urls)

    def _crawl_sitemap_for_documents(self) -> List[str]:
        """Crawl sitemap to discover document URLs."""
        sitemap_variants = ["sitemap.xml", "sitemap_index.xml", "sitemap.php"]
        discovered_urls = set()
        
        for variant in sitemap_variants:
            sitemap_url = urljoin(self.base_url, variant)
            try:
                logger.info(f"Checking sitemap for documents: {sitemap_url}")
                response = self.session.get(sitemap_url, timeout=15)
                
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'xml')
                    urls = {loc.text.strip() for loc in soup.find_all('loc') 
                           if self._is_document_url(loc.text.strip())}
                    discovered_urls.update(urls)
                    logger.info(f"Found {len(urls)} document URLs in {variant}")
                        
            except Exception as e:
                logger.warning(f"Sitemap error ({variant}): {str(e)}")
        
        return list(discovered_urls)

    def _discover_documents_from_pages(self) -> List[str]:
        """Discover document URLs by crawling website pages."""
        document_urls = set()
        pages_to_check = [self.base_url]
        checked_pages = set()
        
        while pages_to_check and len(checked_pages) < 100:  # Limit page crawling
            current_url = pages_to_check.pop(0)
            if current_url in checked_pages:
                continue
                
            checked_pages.add(current_url)
            
            try:
                logger.debug(f"Checking page for documents: {current_url}")
                response = self.session.get(current_url, timeout=10)
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Find all links
                for a_tag in soup.find_all('a', href=True):
                    href = a_tag['href'].strip()
                    if not href:
                        continue
                        
                    full_url = urljoin(current_url, href)
                    
                    # Check if it's a document
                    if self._is_document_url(full_url):
                        # Ensure it's from the same domain
                        if urlparse(full_url).netloc == urlparse(self.base_url).netloc:
                            document_urls.add(full_url)
                    
                    # Add new pages to check (same domain only)
                    elif (not href.startswith('#') and 
                          urlparse(full_url).netloc == urlparse(self.base_url).netloc and
                          full_url not in checked_pages and 
                          full_url not in pages_to_check):
                        pages_to_check.append(full_url)
                
            except Exception as e:
                logger.warning(f"Error checking page {current_url}: {str(e)}")
                continue
        
        logger.info(f"Found {len(document_urls)} document URLs from {len(checked_pages)} pages")
        return list(document_urls)

    def _process_document(self, url: str) -> bool:
        """Process a single document and extract all required data."""
        try:
            logger.info(f"Processing document: {url}")
            
            # Download document
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            if len(response.content) == 0:
                logger.warning(f"Empty document: {url}")
                return False
            
            # Get document type and filename
            doc_type = self._get_document_type(url)
            filename = os.path.basename(urlparse(url).path) or f"document.{doc_type}"
            
            # Process based on document type - Download -> Extract Raw Text
            text_content = ""
            structured_content = {}
            
            if doc_type == 'pdf':
                text_content, structured_content = self.pdf_processor.extract_content(
                    response.content, filename)
            elif doc_type == 'docx':
                text_content, structured_content = self.docx_processor.extract_content(
                    response.content, filename)
            elif doc_type == 'xlsx':
                text_content, structured_content = self.xlsx_processor.extract_content(
                    response.content, filename)
            else:
                logger.warning(f"Unsupported document type: {doc_type}")
                return False
            
            # Skip if no content extracted
            if not text_content or len(text_content.strip()) < 50:
                logger.warning(f"No meaningful content extracted from {filename}")
                return False
            
            # Limit content length for performance
            if len(text_content) > 100000:  # 100k characters max
                text_content = text_content[:100000] + "..."
            
            # Generate title from filename
            title = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title()
            
            # DELEGATE ANALYSIS TO CENTRAL ANALYZER
            # Single, clean call to the new central analyzer after extracting text_content and title
            analysis_results = self.content_analyzer.analyze_content(
                raw_text=text_content,
                title=title,
                metadata={"source_url": url}  # Pass the document URL as metadata
            )
            
            # Simple fallback for content_id
            content_string = f"{url}_{title}"
            content_id = hashlib.md5(content_string.encode('utf-8')).hexdigest()[:12]
            
            # Rebuild document_data dictionary using structured output from analysis_results
            # Map the rich output from analyzer to existing documents.json structure for compatibility
            document_data = {
                "content_id": content_id,
                "url": url,
                "title": analysis_results.get('title', title),
                "summary": analysis_results.get('summary', ''),
                "content": text_content,  # The original raw text from the document
                "search_text": text_content,  # Add search_text field as expected by researcher
                "keywords": analysis_results.get('keywords', []),
                "category": analysis_results.get('primary_category', 'document'),
                "images": [],  # Stays empty for documents
                "internal_links": [],  # Stays empty for documents
                "external_links": [],  # Stays empty for documents
                "structured_content": structured_content,  # This comes from the local extraction helpers
                "metrics": analysis_results.get('readability_metrics', {}),
                "last_updated": datetime.now().isoformat(),
                "processing_info": {
                    "processed_at": datetime.now().isoformat(),
                    "content_length": len(text_content),
                    "document_type": doc_type,
                    "filename": filename,
                    "file_size": len(response.content),
                    "analyzer_version": "centralized_advanced"
                }
            }

            # Add document-specific metrics to the existing metrics from analyzer
            if 'readability_metrics' not in document_data['metrics']:
                document_data['metrics'] = {}
            
            document_data['metrics'].update({
                'file_size_bytes': len(response.content),
                'document_type': doc_type
            })

            # Save individual file (for debugging/backup)
            individual_filename = os.path.join(OUTPUT_DIR, "individual_documents", 
                                             self._safe_filename(url) + ".json")
            os.makedirs(os.path.dirname(individual_filename), exist_ok=True)
            
            with open(individual_filename, 'w', encoding='utf-8') as f:
                json.dump(document_data, f, ensure_ascii=False, indent=2)

            # Store data for consolidated output
            self.documents_data.append(document_data)

            logger.info(f"Successfully processed: {title} ({filename})")
            return True

        except requests.exceptions.Timeout:
            logger.warning(f"Timeout processing {url}")
            return False
        except requests.exceptions.RequestException as e:
            logger.warning(f"Request error for {url}: {str(e)}")
            return False
        except Exception as e:
            logger.error(f"Failed to process {url}: {str(e)}")
            return False

    def _safe_filename(self, url: str) -> str:
        """Generate safe filename from URL."""
        filename = re.sub(r'[^a-zA-Z0-9_-]', '_', 
                         url.replace("https://", "").replace("http://", ""))
        return filename[:100]  # Limit length

    def _save_consolidated_output(self):
        """Save consolidated output in documents.json format compatible with website_researcher.py."""
        # Create the exact structure expected by website_researcher.py
        consolidated_data = {
            "pages": self.documents_data,  # Keep "pages" key for compatibility
            "metadata": {
                "processed_at": datetime.now().isoformat(),
                "last_crawl": datetime.now().isoformat(),  # Keep both for compatibility
                "total_pages": len(self.documents_data),
                "total_urls_found": len(self.documents_data),  # Keep for compatibility
                "version": "2.0",
                "target_website": self.base_url,
                "document_types": self._get_document_type_stats(),
                "processing_summary": {
                    "start_time": datetime.fromtimestamp(self.start_time).isoformat(),
                    "end_time": datetime.now().isoformat(),
                    "duration_seconds": round(time.time() - self.start_time, 2),
                    "categories": self._get_category_stats(),
                    "analyzer_used": "AdvancedContentAnalyzer"
                }
            }
        }
        
        # Save to documents.json
        consolidated_filename = os.path.join(OUTPUT_DIR, "documents.json")
        with open(consolidated_filename, 'w', encoding='utf-8') as f:
            json.dump(consolidated_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Consolidated output saved to {consolidated_filename}")
        logger.info(f"Total documents processed: {len(self.documents_data)}")
        
        # Log statistics
        doc_types = self._get_document_type_stats()
        categories = self._get_category_stats()
        logger.info(f"Document type distribution: {doc_types}")
        logger.info(f"Category distribution: {categories}")

    def _get_category_stats(self) -> Dict[str, int]:
        """Get statistics about document categories."""
        category_counts = {}
        for doc in self.documents_data:
            category = doc.get('category', 'unknown')
            category_counts[category] = category_counts.get(category, 0) + 1
        return category_counts

    def _get_document_type_stats(self) -> Dict[str, int]:
        """Get statistics about document types."""
        type_counts = {}
        for doc in self.documents_data:
            doc_type = doc.get('processing_info', {}).get('document_type', 'unknown')
            type_counts[doc_type] = type_counts.get(doc_type, 0) + 1
        return type_counts

    def run(self):
        """Main processing routine."""
        logger.info("=" * 60)
        logger.info("Starting AI_Gate Document Processing (Centralized Analysis)")
        logger.info("=" * 60)
        
        # Verify accessibility
        accessible, message = self._verify_site_accessibility()
        if not accessible:
            logger.error(f"Access denied: {message}")
            raise SystemExit(f"Fatal Error: {message}")

        logger.info(f"Website access verified: {message}")
        
        try:
            # Discover document URLs
            logger.info("Discovering document URLs...")
            document_urls = self._discover_document_urls()
            
            if not document_urls:
                logger.warning("No document URLs found")
                raise SystemExit("No documents to process")

            logger.info(f"Total documents to process: {len(document_urls)}")
            
            # Process documents
            success_count = 0
            failed_count = 0
            
            for i, url in enumerate(document_urls, 1):
                try:
                    if self._process_document(url):
                        success_count += 1
                    else:
                        failed_count += 1
                        
                    # Progress update
                    if i % 5 == 0 or i == len(document_urls):
                        success_rate = (success_count / i) * 100
                        logger.info(f"Progress: {i}/{len(document_urls)} - "
                                  f"Success: {success_count}, Failed: {failed_count} "
                                  f"({success_rate:.1f}% success rate)")
                        
                    # Small delay to be respectful
                    time.sleep(1)
                    
                except KeyboardInterrupt:
                    logger.warning("Processing interrupted by user")
                    break
                except Exception as e:
                    logger.error(f"Unexpected error processing document {url}: {str(e)}")
                    failed_count += 1
                    continue

            # Save results
            if self.documents_data:
                logger.info("Saving processed data...")
                self._save_consolidated_output()
                logger.info("✅ Document processing completed successfully")
            else:
                logger.warning("⚠️  No documents were successfully processed")

            # Final summary
            logger.info("=" * 60)
            logger.info("PROCESSING SUMMARY")
            logger.info("=" * 60)
            logger.info(f"Total documents attempted: {len(document_urls)}")
            logger.info(f"Successfully processed: {success_count}")
            logger.info(f"Failed to process: {failed_count}")
            logger.info(f"Success rate: {(success_count/len(document_urls)*100):.1f}%")
            logger.info(f"Processing time: {time.time() - self.start_time:.2f} seconds")

        except KeyboardInterrupt:
            logger.warning("Processing interrupted by user")
            # Save whatever data we have collected
            if self.documents_data:
                logger.info("Saving partial results...")
                self._save_consolidated_output()

        except Exception as e:
            logger.critical(f"Critical processing error: {str(e)}", exc_info=True)
            # Try to save partial results
            if self.documents_data:
                logger.info("Attempting to save partial results...")
                try:
                    self._save_consolidated_output()
                except Exception as save_error:
                    logger.error(f"Failed to save partial results: {save_error}")

        finally:
            self._generate_performance_report()

    def _generate_performance_report(self):
        """Generate performance report with detailed statistics."""
        try:
            report = {
                "timestamp": datetime.now().isoformat(),
                "processing_time": round(time.time() - self.start_time, 2),
                "target_url": self.base_url,
                "documents_processed": len(self.documents_data),
                "log_file": os.path.basename(log_file),
                "output_file": os.path.join(OUTPUT_DIR, "documents.json"),
                "compatibility_version": "2.0",
                "analyzer_version": "AdvancedContentAnalyzer (Centralized)",
                "document_types": self._get_document_type_stats(),
                "categories_found": self._get_category_stats(),
                "system_info": {
                    "python_version": os.sys.version,
                    "platform": os.sys.platform
                },
                "processing_notes": [
                    "Output is fully compatible with website_researcher.py",
                    "Individual document files saved to individual_documents/ subdirectory",
                    "Main output saved as documents.json in data directory",
                    "All content analysis performed by centralized AdvancedContentAnalyzer",
                    "High-quality multilingual analysis applied to all documents"
                ]
            }
            
            report_path = os.path.join(LOG_DIR, 
                                     f"performance_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, indent=2, ensure_ascii=False)
                
            logger.info(f"Performance report saved to {report_path}")
            
        except Exception as e:
            logger.error(f"Failed to generate performance report: {str(e)}")

if __name__ == "__main__":
    try:
        processor = DocumentProcessor()
        processor.run()
    except KeyboardInterrupt:
        logger.info("Processing stopped by user")
    except Exception as e:
        logger.critical(f"Application startup error: {str(e)}", exc_info=True)
        raise