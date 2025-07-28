"""
AI_Gate Advanced Multilingual Document Processor - document_processor.py (Version 3.0)

A streamlined document processor that extracts content from local institution documents 
(DOCX, XLSX, PDF) and delegates all NLP analysis to the centralized AdvancedContentAnalyzer.
This version focuses solely on extraction and assembly, maintaining output compatibility 
with the WebsiteResearcher.

Refactoring in v3.0:
- Removed all low-level NLP processing classes (LanguageDetector, AdvancedTextProcessor, etc.)
- Streamlined to focus on extraction and assembly only
- Delegates all content analysis to modules.analyzer.AdvancedContentAnalyzer
- Maintains same output structure for compatibility
- Significantly reduced code complexity and maintenance overhead
"""

import json
import os
import time
import re
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import logging

# Central analyzer import - the only NLP dependency needed
from modules.analyzer import AdvancedContentAnalyzer

# Document processing imports (same as v1.0)
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl import load_workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf
        from pypdf import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Debug configuration
DEBUG = False
if DEBUG:
    logging.basicConfig(level=logging.DEBUG)

# Directories (same as v1.0)
BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = DATA_DIR / "institution_info"
LOG_DIR = BASE_DIR / "logs"

# Create directories
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)

# Logging setup
log_file_path = LOG_DIR / f"document_processor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

root_logger = logging.getLogger()
if root_logger.hasHandlers():
    for handler in root_logger.handlers[:]:
        root_logger.removeHandler(handler)

logging.basicConfig(
    level=logging.DEBUG if DEBUG else logging.INFO,
    format="%(asctime)s - %(levelname)s - %(name)s - %(message)s",
    handlers=[
        logging.FileHandler(log_file_path, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("AI_Gate_StreamlinedDocumentProcessor")

# Log library availability
logger_init = logging.getLogger("AI_Gate_DocumentProcessor_Init")
logger_init.info(f"Streamlined Document Processor v3.0 initializing...")
logger_init.info(f"DOCX Available: {DOCX_AVAILABLE}")
logger_init.info(f"XLSX Available: {XLSX_AVAILABLE}")
logger_init.info(f"PDF Available: {PDF_AVAILABLE}")
logger_init.info(f"PDFPlumber Available: {PDFPLUMBER_AVAILABLE}")


class DocumentProcessor:
    """Streamlined document processor focused on extraction and assembly only."""
    
    def __init__(self, input_directory: str = None):
        self.input_dir = Path(input_directory) if input_directory else DATA_DIR
        self.start_time = time.time()
        self.documents_data: List[Dict] = []
        
        # Single analyzer instance - all NLP work is delegated here
        self.content_analyzer = AdvancedContentAnalyzer()
        
        logger.info(f"Initializing streamlined document processor for directory: {self.input_dir}")
        
        # Verify input directory exists
        if not self.input_dir.exists():
            raise FileNotFoundError(f"Input directory does not exist: {self.input_dir}")

    def _extract_docx_content(self, filepath: Path) -> tuple[str, str, Dict[str, Any]]:
        """Extract content from DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error(f"DOCX processing not available for {filepath}")
            return "", "", {}
        
        try:
            doc = Document(str(filepath))
            
            # Extract title (from document properties or first heading)
            title = ""
            if doc.core_properties.title:
                title = doc.core_properties.title
            else:
                # Try to get title from first paragraph or heading
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        title = paragraph.text.strip()
                        break
            
            if not title:
                title = filepath.stem
            
            # Extract all text content
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            # Extract table content if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(" | ".join(row_text))
            
            content = "\n".join(content_parts)
            
            # Extract metadata
            metadata = {
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "subject": doc.core_properties.subject or "",
                "keywords": doc.core_properties.keywords or ""
            }
            
            return title, content, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {filepath}: {e}")
            return "", "", {}

    def _extract_xlsx_content(self, filepath: Path) -> tuple[str, str, Dict[str, Any]]:
        """Extract content from XLSX file with improved sheet handling."""
        if not XLSX_AVAILABLE:
            logger.error(f"XLSX processing not available for {filepath}")
            return "", "", {}
        
        try:
            workbook = load_workbook(str(filepath), read_only=True, data_only=True)
            
            title = filepath.stem
            content_parts = []
            
            # Process each worksheet with priority to likely content sheets
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                
                # Skip likely metadata sheets
                if any(name.lower() in sheet_name.lower() for name in ['metadata', 'summary', 'info']):
                    continue
                
                # Add sheet name as section header
                content_parts.append(f"\n[Sheet: {sheet_name}]")
                
                # Extract data from first 50 rows (institutional docs rarely need more)
                row_count = 0
                for row in worksheet.iter_rows(values_only=True):
                    row_data = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_data.append(str(cell_value).strip())
                    
                    if row_data and any(cell for cell in row_data if cell):
                        content_parts.append(" | ".join(row_data))
                    
                    row_count += 1
                    if row_count >= 50:
                        break
            
            content = "\n".join(content_parts)
            
            # Basic metadata
            metadata = {
                "sheets": list(workbook.sheetnames),
                "sheet_count": len(workbook.sheetnames)
            }
            
            workbook.close()
            return title, content, metadata
            
        except Exception as e:
            logger.error(f"Error processing XLSX file {filepath}: {e}")
            return "", "", {}

    def _extract_pdf_content(self, filepath: Path) -> tuple[str, str, Dict[str, Any]]:
        """Extract content from PDF file."""
        title, content, metadata = "", "", {}
        
        # Try pdfplumber first (better text extraction)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(str(filepath)) as pdf:
                    title = filepath.stem
                    content_parts = []
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            content_parts.append(f"[Page {page_num}]\n{page_text.strip()}")
                    
                    content = "\n\n".join(content_parts)
                    
                    metadata = {
                        "pages": len(pdf.pages),
                        "extractor": "pdfplumber"
                    }
                    
                    # Try to get PDF metadata
                    if hasattr(pdf, 'metadata') and pdf.metadata:
                        if pdf.metadata.get('Title'):
                            title = pdf.metadata['Title']
                        metadata.update({
                            "author": pdf.metadata.get('Author', ''),
                            "subject": pdf.metadata.get('Subject', ''),
                            "creator": pdf.metadata.get('Creator', '')
                        })
                
                if content.strip():
                    return title, content, metadata
                    
            except Exception as e:
                logger.warning(f"pdfplumber failed for {filepath}: {e}")
        
        # Fallback to PyPDF2/pypdf
        if PDF_AVAILABLE:
            try:
                with open(filepath, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    
                    title = filepath.stem
                    content_parts = []
                    
                    # Extract text from all pages
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                content_parts.append(f"[Page {page_num}]\n{page_text.strip()}")
                        except Exception as e:
                            logger.warning(f"Error extracting page {page_num} from {filepath}: {e}")
                    
                    content = "\n\n".join(content_parts)
                    
                    metadata = {
                        "pages": len(pdf_reader.pages),
                        "extractor": "PyPDF2/pypdf"
                    }
                    
                    # Try to get PDF metadata
                    if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                        if pdf_reader.metadata.get('/Title'):
                            title = pdf_reader.metadata['/Title']
                        metadata.update({
                            "author": pdf_reader.metadata.get('/Author', ''),
                            "subject": pdf_reader.metadata.get('/Subject', ''),
                            "creator": pdf_reader.metadata.get('/Creator', '')
                        })
                
                return title, content, metadata
                
            except Exception as e:
                logger.error(f"Error processing PDF file {filepath}: {e}")
        
        return title, content, metadata

    def _process_document(self, filepath: Path) -> bool:
        """Process a single document file using centralized content analysis."""
        try:
            logger.info(f"Processing document: {filepath}")
            
            # Step 1: Determine file type and extract raw content
            suffix = filepath.suffix.lower()
            title, content, metadata = "", "", {}
            
            if suffix == '.docx':
                title, content, metadata = self._extract_docx_content(filepath)
            elif suffix == '.xlsx':
                title, content, metadata = self._extract_xlsx_content(filepath)
            elif suffix == '.pdf':
                title, content, metadata = self._extract_pdf_content(filepath)
            else:
                logger.warning(f"Unsupported file type: {suffix} for {filepath}")
                return False
            
            if not content or not content.strip():
                logger.warning(f"No content extracted from {filepath}")
                return False
            
            # Step 2: Clean and limit content
            cleaned_content = re.sub(r'\s+', ' ', content).strip()
            if len(cleaned_content) > 75000:
                cleaned_content = cleaned_content[:75000] + "..."
            
            # Step 3: Delegate ALL analysis to the centralized analyzer
            analysis_results = self.content_analyzer.analyze_content(
                raw_text=cleaned_content,
                title=title,
                metadata={"source_file": str(filepath), "file_metadata": metadata}
            )
            
            # Step 4: Generate content ID using analyzer's method (or local fallback)
            try:
                content_id = self.content_analyzer.generate_content_id(str(filepath), title)
            except AttributeError:
                # Fallback if method doesn't exist in analyzer
                content_string = f"{filepath}_{title}_{datetime.now().date()}"
                content_id = hashlib.md5(content_string.encode('utf-8')).hexdigest()[:12]
            
            # Step 5: Assemble the final document data using analyzer results
            document_data = {
                "content_id": content_id,
                "url": str(filepath),  # Using filepath as URL equivalent
                "title": analysis_results.get('title', title),
                "summary": analysis_results.get('summary', ''),
                "content": cleaned_content,
                "search_text": cleaned_content,  # Same as content for documents
                "keywords": analysis_results.get('keywords', []),
                "category": analysis_results.get('primary_category', 'general'),
                "images": [],  # Documents don't have image URLs
                "internal_links": [],  # Not applicable for documents
                "external_links": [],  # Not applicable for documents
                "structured_content": analysis_results.get('structured_content', {}),
                "metrics": analysis_results.get('readability_metrics', {}),
                "last_updated": datetime.now().isoformat(),
                "processing_info": {
                    "processed_at": datetime.now().isoformat(),
                    "content_length": len(cleaned_content),
                    "links_found": 0,  # N/A for documents
                    "images_found": 0,  # N/A for documents
                    "javascript_rendered": False,  # N/A for documents
                    "file_type": suffix,
                    "file_size_bytes": filepath.stat().st_size,
                    "metadata": metadata,
                    "analyzer_version": "3.0",
                    "analysis_features_used": list(analysis_results.keys())
                }
            }
            
            self.documents_data.append(document_data)
            logger.info(f"Successfully processed: {document_data['title']} ({filepath.name})")
            return True
            
        except Exception as e:
            logger.error(f"Failed to process document {filepath}: {e}")
            return False

    def _find_supported_documents(self) -> List[Path]:
        """Find all supported document files in the input directory."""
        supported_extensions = ['.docx', '.xlsx', '.pdf']
        documents = []
        
        for ext in supported_extensions:
            pattern = f"*{ext}"
            found_files = list(self.input_dir.glob(pattern))
            documents.extend(found_files)
            logger.info(f"Found {len(found_files)} {ext.upper()} files")
        
        # Also check subdirectories
        for ext in supported_extensions:
            pattern = f"**/*{ext}"
            found_files = list(self.input_dir.glob(pattern))
            # Remove duplicates from direct directory scan
            new_files = [f for f in found_files if f not in documents]
            documents.extend(new_files)
            if new_files:
                logger.info(f"Found {len(new_files)} additional {ext.upper()} files in subdirectories")
        
        return sorted(documents)

    def _save_consolidated_output(self):
        """Save consolidated output to JSON file."""
        if not self.documents_data:
            logger.warning("No document data to save.")
            return
        
        # Create metadata matching website_scraper.py format
        metadata = {
            "processed_at": datetime.now().isoformat(),
            "last_crawl": datetime.now().isoformat(),  # Using same field name for compatibility
            "total_pages": len(self.documents_data),  # Using same field name for compatibility
            "total_urls_found": len(self.documents_data),  # Using same field name for compatibility
            "version": "3.0",  # Updated version number
            "target_website": str(self.input_dir),  # Using input directory as "target"
            "javascript_support_used": False,  # N/A for documents
            "ssl_verification_status": "N/A",  # N/A for documents
            "processing_summary": {
                "start_time": datetime.fromtimestamp(self.start_time).isoformat(),
                "end_time": datetime.now().isoformat(),
                "duration_seconds": round(time.time() - self.start_time, 2),
                "categories": self._get_category_stats(),
                "javascript_rendered_pages_count": 0,  # N/A for documents
                "document_types_processed": self._get_file_type_stats(),
                "analyzer_used": "modules.analyzer.AdvancedContentAnalyzer"
            }
        }
        
        # Create consolidated data structure matching website_scraper.py
        consolidated_data = {
            "pages": self.documents_data,  # Using same field name for compatibility
            "metadata": metadata
        }
        
        # Save to output file
        output_path = OUTPUT_DIR / "ins_info.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(consolidated_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Consolidated output saved to {output_path} ({len(self.documents_data)} documents)")
        logger.info(f"Category distribution: {self._get_category_stats()}")
        logger.info(f"File type distribution: {self._get_file_type_stats()}")

    def _get_category_stats(self) -> Dict[str, int]:
        """Get category distribution statistics."""
        counts: Dict[str, int] = {}
        for doc in self.documents_data:
            category = doc.get('category', 'unknown')
            counts[category] = counts.get(category, 0) + 1
        return counts

    def _get_file_type_stats(self) -> Dict[str, int]:
        """Get file type distribution statistics."""
        counts: Dict[str, int] = {}
        for doc in self.documents_data:
            file_type = doc.get('processing_info', {}).get('file_type', 'unknown')
            counts[file_type] = counts.get(file_type, 0) + 1
        return counts

    def run(self):
        """Main processing routine."""
        logger.info("=" * 60)
        logger.info("Starting AI_Gate Streamlined Document Processing (v3.0)")
        logger.info("=" * 60)
        
        try:
            # Find all supported documents
            documents = self._find_supported_documents()
            
            if not documents:
                logger.warning(f"No supported documents found in {self.input_dir}")
                logger.info("Supported formats: DOCX, XLSX, PDF")
                return
            
            logger.info(f"Found {len(documents)} documents to process")
            
            # Process each document
            success_count = 0
            failed_count = 0
            
            for i, doc_path in enumerate(documents, 1):
                try:
                    if self._process_document(doc_path):
                        success_count += 1
                    else:
                        failed_count += 1
                    
                    # Progress update
                    if i % 5 == 0 or i == len(documents):
                        rate = (success_count / i) * 100 if i > 0 else 0
                        logger.info(f"Progress: {i}/{len(documents)} - Success: {success_count}, Failed: {failed_count} ({rate:.1f}%)")
                
                except KeyboardInterrupt:
                    logger.warning("Processing interrupted by user")
                    break
                except Exception as e:
                    logger.error(f"Unexpected error processing {doc_path}: {e}")
                    failed_count += 1
                    continue
            
            # Save results
            if self.documents_data:
                logger.info("Saving processed data...")
                self._save_consolidated_output()
                logger.info("✅ Document processing completed successfully.")
            else:
                logger.warning("⚠️ No documents were successfully processed.")
            
            # Summary
            logger.info("=" * 60)
            logger.info("DOCUMENT PROCESSING SUMMARY")
            logger.info("=" * 60)
            logger.info(f"Total documents found: {len(documents)}")
            logger.info(f"Successfully processed: {success_count}")
            logger.info(f"Failed: {failed_count}")
            if len(documents) > 0:
                logger.info(f"Success rate: {(success_count/len(documents)*100):.1f}%")
            logger.info(f"Total processing time: {time.time() - self.start_time:.2f}s")
            
        except KeyboardInterrupt:
            logger.warning("Processing interrupted by user")
        except Exception as e:
            logger.critical(f"Critical error: {e}", exc_info=True)
        finally:
            # Save partial results if any processing occurred
            if self.documents_data and (success_count > 0 or failed_count > 0):
                logger.info("Ensuring results are saved...")
                try:
                    self._save_consolidated_output()
                except Exception as save_error:
                    logger.error(f"Failed to save results: {save_error}")
            
            self._generate_performance_report()

    def _generate_performance_report(self):
        """Generate performance report."""
        try:
            report = {
                "timestamp": datetime.now().isoformat(),
                "processing_time_seconds": round(time.time() - self.start_time, 2),
                "input_directory": str(self.input_dir),
                "documents_processed": len(self.documents_data),
                "output_file": str(OUTPUT_DIR / "ins_info.json"),
                "version": "3.0",  # Updated version number
                "categories_found": self._get_category_stats(),
                "file_types_processed": self._get_file_type_stats(),
                "system_info": {
                    "docx_available": DOCX_AVAILABLE,
                    "xlsx_available": XLSX_AVAILABLE,
                    "pdf_available": PDF_AVAILABLE,
                    "pdfplumber_available": PDFPLUMBER_AVAILABLE,
                    "analyzer_module": "modules.analyzer.AdvancedContentAnalyzer"
                },
                "processing_notes": [
                    "Streamlined extraction and assembly processor",
                    "All NLP analysis delegated to modules.analyzer.AdvancedContentAnalyzer",
                    "Output structure compatible with WebsiteResearcher",
                    "Documents processed from local files",
                    "Main output: ins_info.json",
                    "Maintains same JSON structure as pages.json"
                ]
            }
            
            report_path = LOG_DIR / f"performance_report_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Performance report saved to {report_path}")
            
        except Exception as e:
            logger.error(f"Failed to generate performance report: {e}")


if __name__ == "__main__":
    try:
        # Check if required libraries are available
        missing_libs = []
        if not DOCX_AVAILABLE:
            missing_libs.append("python-docx (for DOCX files)")
        if not XLSX_AVAILABLE:
            missing_libs.append("openpyxl (for XLSX files)")
        if not PDF_AVAILABLE:
            missing_libs.append("PyPDF2 or pypdf (for PDF files)")
        
        if missing_libs:
            logger.warning("Missing optional libraries:")
            for lib in missing_libs:
                logger.warning(f"  - {lib}")
            logger.warning("Some file types may not be processed.")
        
        # Create and run processor
        processor = DocumentProcessor()
        processor.run()
        
    except KeyboardInterrupt:
        logger.info("Processing stopped by user")
    except Exception as e:
        logger.critical(f"Application startup error: {e}", exc_info=True)